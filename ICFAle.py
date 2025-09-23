# streamlit_app.py
import streamlit as st
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import zipfile
import re

st.set_page_config(page_title="Generador DOCX por protocolo", layout="wide")

st.title("Generador automático de Consentimientos (python-docx + Streamlit)")

st.markdown("""
Subí tu **modelo.docx** (plantilla) con placeholders (ej: `{{INVESTIGADOR}}`) y el **datos.xlsx** con una fila por investigador.
El nombre del archivo resultante incluirá el número de protocolo y el nombre del investigador.
""")

uploaded_docx = st.file_uploader("Subí el modelo (.docx)", type=["docx"])
uploaded_xlsx = st.file_uploader("Subí el Excel (.xlsx)", type=["xlsx"])

# Texto a buscar / reemplazar para casos especiales
texto_anticonceptivo_original = (
    "El médico del estudio discutirá con usted qué método anticonceptivo se considera adecuado. "
    "El patrocinador y/o el investigador del estudio garantizarán su acceso al método anticonceptivo "
    "acordado y necesario para su participación en este estudio"
)

texto_ba_reemplazo = (
    "El médico del estudio discutirá con usted qué métodos anticonceptivos se consideran adecuados. "
    "El Patrocinador y/o el médico del estudio garantizará su acceso a este método anticonceptivo "
    "acordado y necesario para su participación en el ensayo. El costo de los métodos anticonceptivos "
    "seleccionados correrá a cargo del Patrocinador."
)

# Funciones auxiliares
def remove_paragraph(paragraph):
    # Elimina un párrafo del documento (manipulando XML)
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def replace_text_in_runs(paragraph, old, new):
    """
    Reemplaza texto buscando en cada run individual. Si el placeholder
    está dividido en varias runs, puede que no lo encuentre.
    """
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

def replace_text_in_doc(doc, replacements):
    # Reemplazo en paragraphs
    for p in doc.paragraphs:
        for old, new in replacements.items():
            # reemplazo por runs (simple y rápido)
            replace_text_in_runs(p, old, new)
        # también chequeo si algún placeholder quedó en el texto completo del párrafo:
        # (esto cubre casos donde placeholder no fue encontrado por run)
        fulltext = p.text
        for old, new in replacements.items():
            if old in fulltext:
                # si queda, borramos runs y escribimos nuevo texto (puede cambiar formato)
                for r in p.runs:
                    r.text = ""
                p.add_run(fulltext.replace(old, new))
    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        replace_text_in_runs(p, old, new)
                    fulltext = p.text
                    for old, new in replacements.items():
                        if old in fulltext:
                            for r in p.runs:
                                r.text = ""
                            p.add_run(fulltext.replace(old, new))

def find_paragraphs_containing(doc, snippet):
    res = []
    for p in doc.paragraphs:
        if snippet.lower() in p.text.lower():
            res.append(p)
    # también podemos buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if snippet.lower() in p.text.lower():
                        res.append(p)
    return res

def process_row_and_generate_doc(template_bytes, row):
    # row: pandas Series
    doc = Document(io.BytesIO(template_bytes))

    # Mapeo de placeholders a valores
    replacements = {
        "{{NUM_PROTOCOLO}}": str(row.get("Numero de protocolo", "")).strip(),
        "{{TITULO_ESTUDIO}}": str(row.get("Titulo del Estudio", "")).strip(),
        "{{PATROCINADOR}}": str(row.get("Patrocinador", "")).strip(),
        "{{INVESTIGADOR}}": str(row.get("Investigador", "")).strip(),
        "{{INSTITUCION}}": str(row.get("Institucion", "")).strip(),
        "{{DIRECCION}}": str(row.get("Direccion", "")).strip(),
        "{{CARGO}}": str(row.get("Cargo", "")).strip(),
        "{{PROVINCIA}}": str(row.get("provincia", "")).strip(),
        "{{COMITE}}": str(row.get("comite", "")).strip(),
        "{{SUBINVESTIGADOR}}": str(row.get("subinvestigador", "")).strip(),
        "{{TELEFONO_24HS}}": str(row.get("TELEFONO_24HS", "")).strip(),
    }

    # Reemplazamos placeholders
    replace_text_in_doc(doc, replacements)

    # Lógica por provincia
    prov = str(row.get("provincia", "")).strip().lower()

    # 1) Si cordoba -> eliminar el párrafo que contiene texto_anticonceptivo_original
    if prov == "cordoba":
        paras = find_paragraphs_containing(doc, texto_anticonceptivo_original)
        for p in paras:
            try:
                remove_paragraph(p)
            except Exception:
                pass
        # también eliminamos cualquier párrafo BA si existiera
        paras_ba = find_paragraphs_containing(doc, "Requerido para centros de la provincia de Buenos Aires")
        for p in paras_ba:
            try:
                remove_paragraph(p)
            except Exception:
                pass

    # 2) Si Buenos Aires -> reemplazar con texto_ba_reemplazo
    elif prov.replace(" ", "") in ("buenosaires", "buenos aires", "buenosaires"):
        paras = find_paragraphs_containing(doc, texto_anticonceptivo_original)
        if paras:
            for p in paras:
                # borramos runs y ponemos el texto nuevo
                for r in p.runs:
                    r.text = ""
                p.add_run(texto_ba_reemplazo)
        else:
            # si no se encontró, intentamos buscar el párrafo "Requerido para centros..."
            paras_ba = find_paragraphs_containing(doc, "Requerido para centros de la provincia de Buenos Aires")
            for p in paras_ba:
                for r in p.runs:
                    r.text = ""
                p.add_run(texto_ba_reemplazo)

    # Guardamos a bytes y retornamos
    out_io = io.BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    return out_io

# Botón para ejecutar
if uploaded_docx and uploaded_xlsx:
    try:
        df = pd.read_excel(uploaded_xlsx, engine="openpyxl")
    except Exception as e:
        st.error(f"Error leyendo el Excel: {e}")
        st.stop()

    # Normalizamos nombres de columnas (por si acaso)
    df_columns_lower = {c.lower(): c for c in df.columns}
    # chequeo rápido de columnas esperadas
    expected = ["numero de protocolo", "titulo del estudio", "patrocinador", "investigador",
                "institucion", "direccion", "cargo", "provincia", "comite", "subinvestigador", "telefono_24hs"]
    # No descartamos si faltan, pero avisamos
    missing = [e for e in expected if e not in df_columns_lower]
    if missing:
        st.warning("Verifica que tu Excel contenga (al menos) estas columnas (case-insensitive):\n" + ", ".join(expected))
        st.info("Continuaré intentando con las columnas presentes (usando nombres exactos).")

    # convertimos template a bytes
    template_bytes = uploaded_docx.read()

    # Creamos ZIP en memoria
    zip_io = io.BytesIO()
    with zipfile.ZipFile(zip_io, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        # iteramos filas
        for idx, row in df.iterrows():
            # procesamos la fila
            try:
                doc_io = process_row_and_generate_doc(template_bytes, row)
            except Exception as e:
                st.error(f"Error procesando fila {idx}: {e}")
                continue

            # Generar nombre de archivo: <numProtocolo> - <Investigador>.docx
            num_prot = str(row.get("Numero de protocolo", "")).strip()
            inv = str(row.get("Investigador", "")).strip()
            safe_num = re.sub(r'[\\/*?:"<>|]', "_", num_prot)[:120]
            safe_inv = re.sub(r'[\\/*?:"<>|]', "_", inv)[:120]
            filename = f"{safe_num} - {safe_inv}.docx" if safe_num or safe_inv else f"doc_{idx}.docx"

            zf.writestr(filename, doc_io.getvalue())

    zip_io.seek(0)
    st.success("Generados todos los documentos.")
    st.download_button("Descargar ZIP con todos los DOCX", data=zip_io.getvalue(),
                       file_name="consentimientos_generados.zip", mime="application/zip")
else:
    st.info("Subí el modelo .docx y el .xlsx para comenzar.")
